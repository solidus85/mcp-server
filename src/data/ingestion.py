import asyncio
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import mimetypes
import logging
from dataclasses import dataclass
from enum import Enum

from ..utils import setup_logging, chunk_text, hash_text, generate_id
from .transformers import DataTransformer, DataCleaner


logger = setup_logging("DataIngestion")


class DocumentType(str, Enum):
    """Supported document types"""
    TEXT = "text"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"
    CSV = "csv"
    XML = "xml"
    DOCX = "docx"


@dataclass
class Document:
    """Represents a document for ingestion"""
    id: str
    content: str
    metadata: Dict[str, Any]
    type: DocumentType
    source: Optional[str] = None
    chunks: Optional[List[str]] = None


class DocumentProcessor:
    """Process various document types for ingestion"""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        clean_html: bool = True,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.clean_html = clean_html
        self.cleaner = DataCleaner()
        self.transformer = DataTransformer()
    
    async def process_file(self, file_path: Path) -> Document:
        """Process a file and return a Document"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        doc_type = self._get_document_type(file_path, mime_type)
        
        # Read and process content
        content = await self._read_file(file_path, doc_type)
        
        # Create document
        doc_id = generate_id("doc")
        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "mime_type": mime_type,
            "doc_type": doc_type.value,
        }
        
        return Document(
            id=doc_id,
            content=content,
            metadata=metadata,
            type=doc_type,
            source=str(file_path),
        )
    
    async def process_text(
        self,
        text: str,
        doc_type: DocumentType = DocumentType.TEXT,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Process raw text into a Document"""
        # Clean text if needed
        if doc_type == DocumentType.HTML and self.clean_html:
            text = self.cleaner.sanitize_html(text)
        
        # Create document
        doc_id = generate_id("doc")
        doc_metadata = metadata or {}
        doc_metadata.update({
            "doc_type": doc_type.value,
            "content_hash": hash_text(text),
            "content_length": len(text),
        })
        
        return Document(
            id=doc_id,
            content=text,
            metadata=doc_metadata,
            type=doc_type,
        )
    
    def chunk_document(self, document: Document) -> Document:
        """Split document into chunks"""
        chunks = chunk_text(
            document.content,
            chunk_size=self.chunk_size,
            overlap=self.chunk_overlap,
        )
        
        document.chunks = chunks
        document.metadata["num_chunks"] = len(chunks)
        
        return document
    
    async def _read_file(self, file_path: Path, doc_type: DocumentType) -> str:
        """Read file content based on type"""
        if doc_type == DocumentType.PDF:
            return await self._read_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return await self._read_docx(file_path)
        elif doc_type in [DocumentType.JSON, DocumentType.CSV, DocumentType.XML]:
            content = file_path.read_text(encoding="utf-8")
            # Optionally convert to plain text
            if doc_type == DocumentType.JSON:
                import json
                data = json.loads(content)
                return json.dumps(data, indent=2)
            return content
        else:
            # Default to text reading
            return file_path.read_text(encoding="utf-8")
    
    async def _read_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            text = []
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            
            return "\n\n".join(text)
        except ImportError:
            logger.warning("PyPDF2 not installed, treating PDF as binary")
            return f"[PDF file: {file_path.name}]"
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return f"[Error reading PDF: {file_path.name}]"
    
    async def _read_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            return "\n\n".join(text)
        except ImportError:
            logger.warning("python-docx not installed, treating DOCX as binary")
            return f"[DOCX file: {file_path.name}]"
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return f"[Error reading DOCX: {file_path.name}]"
    
    def _get_document_type(self, file_path: Path, mime_type: Optional[str]) -> DocumentType:
        """Determine document type from file extension and MIME type"""
        ext = file_path.suffix.lower()
        
        if ext == ".pdf":
            return DocumentType.PDF
        elif ext == ".docx":
            return DocumentType.DOCX
        elif ext in [".html", ".htm"]:
            return DocumentType.HTML
        elif ext in [".md", ".markdown"]:
            return DocumentType.MARKDOWN
        elif ext == ".json":
            return DocumentType.JSON
        elif ext == ".csv":
            return DocumentType.CSV
        elif ext == ".xml":
            return DocumentType.XML
        else:
            return DocumentType.TEXT


class DocumentIngestionPipeline:
    """Complete pipeline for document ingestion"""
    
    def __init__(
        self,
        processor: Optional[DocumentProcessor] = None,
        enable_chunking: bool = True,
        enable_deduplication: bool = True,
    ):
        self.processor = processor or DocumentProcessor()
        self.enable_chunking = enable_chunking
        self.enable_deduplication = enable_deduplication
        self.processed_hashes = set()
    
    async def ingest_files(
        self,
        file_paths: List[Path],
        batch_size: int = 10,
    ) -> List[Document]:
        """Ingest multiple files"""
        documents = []
        
        # Process in batches
        for i in range(0, len(file_paths), batch_size):
            batch = file_paths[i:i+batch_size]
            
            # Process files concurrently
            tasks = [self.processor.process_file(fp) for fp in batch]
            batch_docs = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Filter out errors and process successful documents
            for doc in batch_docs:
                if isinstance(doc, Exception):
                    logger.error(f"Error processing file: {doc}")
                else:
                    processed_doc = await self._process_document(doc)
                    if processed_doc:
                        documents.append(processed_doc)
        
        logger.info(f"Ingested {len(documents)} documents from {len(file_paths)} files")
        return documents
    
    async def ingest_texts(
        self,
        texts: List[str],
        metadatas: Optional[List[Dict[str, Any]]] = None,
    ) -> List[Document]:
        """Ingest multiple text documents"""
        documents = []
        metadatas = metadatas or [{}] * len(texts)
        
        for text, metadata in zip(texts, metadatas):
            doc = await self.processor.process_text(text, metadata=metadata)
            processed_doc = await self._process_document(doc)
            if processed_doc:
                documents.append(processed_doc)
        
        logger.info(f"Ingested {len(documents)} text documents")
        return documents
    
    async def ingest_directory(
        self,
        directory: Path,
        pattern: str = "*",
        recursive: bool = True,
    ) -> List[Document]:
        """Ingest all matching files from a directory"""
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory: {directory}")
        
        # Find matching files
        if recursive:
            file_paths = list(directory.rglob(pattern))
        else:
            file_paths = list(directory.glob(pattern))
        
        # Filter to only files
        file_paths = [fp for fp in file_paths if fp.is_file()]
        
        logger.info(f"Found {len(file_paths)} files matching pattern '{pattern}'")
        
        return await self.ingest_files(file_paths)
    
    async def _process_document(self, document: Document) -> Optional[Document]:
        """Process a single document through the pipeline"""
        # Check for duplicates
        if self.enable_deduplication:
            content_hash = hash_text(document.content)
            if content_hash in self.processed_hashes:
                logger.debug(f"Skipping duplicate document: {document.id}")
                return None
            self.processed_hashes.add(content_hash)
        
        # Chunk document if enabled
        if self.enable_chunking:
            document = self.processor.chunk_document(document)
        
        # Add processing metadata
        document.metadata["processed"] = True
        document.metadata["pipeline_version"] = "1.0"
        
        return document
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get pipeline statistics"""
        return {
            "processed_documents": len(self.processed_hashes),
            "chunking_enabled": self.enable_chunking,
            "deduplication_enabled": self.enable_deduplication,
            "chunk_size": self.processor.chunk_size,
            "chunk_overlap": self.processor.chunk_overlap,
        }